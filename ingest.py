# When querying the vector store, we need to ensure that we properly query for either the job descriptions or the schema/example JSON documents.
# To do this, we add a filter argument to the similarity_search method to specify the source type.
# For job desctriptions, use vectorstore.similarity_search(question, k=5, filter={"source": "job_description"})
# For Schema/example JSON documents, use vectorstore.similarity_search(question, k=3, filter={"source": "resume_schema"})
import os
import json
import pdfplumber
from docx import Document as DocxDocument
from ollama import Client
import numpy as np
import faiss
from langchain_core.documents import Document as LangChainDocument
from langchain_core.embeddings import Embeddings
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from typing import List, Dict, Any

# Ollama client
client = Client()

class OllamaEmbeddings(Embeddings):
    """Custom embedding function wrapper for LangChain"""
    def __init__(self, model="mxbai-embed-large"):
        self.model = model
        self.client = Client()
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        embeddings = []
        for i, text in enumerate(texts):
            if i % 10 == 0:
                print(f"  Embedding document {i+1}/{len(texts)}...")
            response = self.client.embeddings(
                model=self.model,
                prompt=text
            )
            embeddings.append(response["embedding"])
        return embeddings
    
    def embed_query(self, text: str) -> List[float]:
        response = self.client.embeddings(
            model=self.model,
            prompt=text
        )
        return response["embedding"]


def load_pdf_text(path: str) -> str:
    """Extract text from PDF"""
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def load_docx_text(path: str) -> str:
    """Extract text from DOCX"""
    doc = DocxDocument(path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return text


def chunk_job_description(job: Dict[str, Any], job_id: int) -> List[LangChainDocument]:
    """
    Intelligently chunk a job description into semantic sections.
    Each chunk preserves metadata for better retrieval.
    Tailored for your JSON structure with fields like:
    - position_title, company_name, location
    - position_summary, company_description
    - primary_responsibilities, requirements
    - salary_min, salary_max, job_type
    """
    chunks = []
    
    # Extract standard fields
    job_title = job.get("position_title", "Unknown Position")
    company = job.get("company_name", "Unknown Company")
    location = job.get("location", "")
    job_type = job.get("job_type", "")
    job_id_from_data = job.get("id", job_id)
    
    # Salary info for metadata
    salary_info = None
    if job.get("salary_min") and job.get("salary_max"):
        pay_period = job.get("pay_period", "Annual")
        salary_info = f"${job['salary_min']:,} - ${job['salary_max']:,} {pay_period}"
    
    # 1. OVERVIEW CHUNK (Company + Position Summary)
    overview_parts = [
        f"Position: {job_title}",
        f"Company: {company}"
    ]
    
    if location:
        overview_parts.append(f"Location: {location}")
    if job_type:
        overview_parts.append(f"Job Type: {job_type}")
    
    # Add company description
    if job.get("company_description"):
        overview_parts.append(f"\nAbout the Company:\n{job['company_description']}")
    
    # Add position summary
    if job.get("position_summary"):
        overview_parts.append(f"\nPosition Summary:\n{job['position_summary']}")
    
    chunks.append(LangChainDocument(
        page_content="\n".join(overview_parts),
        metadata={
            "job_id": job_id_from_data,
            "job_title": job_title,
            "company": company,
            "location": location,
            "job_type": job_type,
            "salary": salary_info,
            "section": "overview",
            "source": "job_description"
        }
    ))
    
    # 2. RESPONSIBILITIES CHUNK
    if job.get("primary_responsibilities"):
        responsibilities = job["primary_responsibilities"]
        if isinstance(responsibilities, list):
            content = "\n".join(f"• {item}" for item in responsibilities)
        else:
            content = str(responsibilities)
        
        chunks.append(LangChainDocument(
            page_content=f"Primary Responsibilities:\n{content}",
            metadata={
                "job_id": job_id_from_data,
                "job_title": job_title,
                "company": company,
                "location": location,
                "job_type": job_type,
                "salary": salary_info,
                "section": "responsibilities",
                "source": "job_description"
            }
        ))
    
    # 3. REQUIREMENTS CHUNK
    if job.get("requirements"):
        requirements = job["requirements"]
        if isinstance(requirements, list):
            content = "\n".join(f"• {item}" for item in requirements)
        else:
            content = str(requirements)
        
        chunks.append(LangChainDocument(
            page_content=f"Requirements:\n{content}",
            metadata={
                "job_id": job_id_from_data,
                "job_title": job_title,
                "company": company,
                "location": location,
                "job_type": job_type,
                "salary": salary_info,
                "section": "requirements",
                "source": "job_description"
            }
        ))
    
    # 4. COMPENSATION CHUNK (if you want salary searchable)
    if salary_info and job.get("pay_period"):
        comp_parts = [f"Compensation: {salary_info}"]
        
        # Add any benefits or relocation info if present
        if job.get("benefits"):
            benefits = job["benefits"]
            if isinstance(benefits, list):
                comp_parts.append("Benefits:\n" + "\n".join(f"• {b}" for b in benefits))
            else:
                comp_parts.append(f"Benefits: {benefits}")
        
        if job.get("relocation_required") is not None:
            reloc = "Required" if job["relocation_required"] else "Not Required"
            comp_parts.append(f"Relocation: {reloc}")
        
        chunks.append(LangChainDocument(
            page_content="\n".join(comp_parts),
            metadata={
                "job_id": job_id_from_data,
                "job_title": job_title,
                "company": company,
                "location": location,
                "job_type": job_type,
                "salary": salary_info,
                "section": "compensation",
                "source": "job_description"
            }
        ))
    
    return chunks


def load_job_descriptions(json_path: str) -> List[LangChainDocument]:
    """
    Load and chunk job descriptions from JSON file.
    Returns a list of LangChain Documents with metadata.
    """
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    # Normalize to list of dicts
    if isinstance(data, dict):
        jobs = [data]
    elif isinstance(data, list):
        jobs = data
    else:
        raise ValueError("JSON must be a dict or list of dicts")
    
    all_chunks = []
    for i, job in enumerate(jobs):
        if not isinstance(job, dict):
            print(f"Warning: Job {i} is not a dict, skipping")
            continue
        
        chunks = chunk_job_description(job, job_id=i)
        all_chunks.extend(chunks)
        
        if (i + 1) % 10 == 0:
            print(f"  Processed {i + 1}/{len(jobs)} jobs...")
    
    print(f"✓ Loaded {len(jobs)} jobs → {len(all_chunks)} chunks")
    return all_chunks

def load_example_json_documents(folder: str) -> List[LangChainDocument]:
    """
    Load schema/example JSON files as single documents.
    Each file becomes one retrievable document with metadata tags.
    Currently unused
    """
    docs = []
    for filename in os.listdir(folder):
        if not filename.endswith(".json"):
            continue

        path = os.path.join(folder, filename)
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        docs.append(
            LangChainDocument(
                page_content=content,
                metadata={
                    "source": "resume_schema",
                    "filename": filename,
                    "doc_type": "schema_or_example"
                }
            )
        )

    print(f"✓ Loaded {len(docs)} schema/example JSON files")
    return docs


def build_faiss_vectorstore(documents: List[LangChainDocument]) -> FAISS:
    """
    Build FAISS vector store from documents.
    """
    print(f"\nGenerating embeddings for {len(documents)} documents...")
    embedding_function = OllamaEmbeddings()
    
    # Use LangChain's from_documents method (handles embedding automatically)
    vectorstore = FAISS.from_documents(
        documents=documents,
        embedding=embedding_function
    )
    
    return vectorstore


def save_vectorstore(vectorstore: FAISS, path: str = "vectorstore_jobs"):
    """Save vector store to disk"""
    os.makedirs(path, exist_ok=True)
    vectorstore.save_local(path)
    print(f"✓ Vector store saved to {path}/")


def search_jobs(vectorstore: FAISS, query: str, k: int = 5) -> List[LangChainDocument]:
    """
    Search for relevant job description chunks.
    Returns documents with metadata intact.
    """
    results = vectorstore.similarity_search(query, k=k)
    return results


# --- MAIN ---
if __name__ == "__main__":
    print("=== Job Description Vector Store Builder ===\n")
    
    # Path to your JSON file
    json_path = "data/job_description/job_descriptions_100.json"
    
    if not os.path.exists(json_path):
        print(f"Error: {json_path} not found!")
        exit(1)
    
    # Load and chunk job descriptions
    print("Loading job descriptions...")
    documents = load_job_descriptions(json_path)
    
    if not documents:
        print("Error: No documents were created. Check your JSON structure.")
        exit(1)
    
    # Build vector store
    vectorstore = build_faiss_vectorstore(documents)
    
    # Save to disk
    save_vectorstore(vectorstore)
    
    print(f"\n✅ Successfully indexed {len(documents)} document chunks!")
    
    # Test search
    print("\n=== Testing Search ===")
    test_queries = [
        "Python backend development experience",
        "leadership and team management",
        "machine learning and AI skills"
    ]
    
    for query in test_queries:
        print(f"\n🔍 Query: '{query}'")
        results = search_jobs(vectorstore, query, k=3)
        for i, doc in enumerate(results, 1):
            print(f"\n  Result {i}:")
            print(f"  Job: {doc.metadata.get('job_title')} at {doc.metadata.get('company')}")
            print(f"  Section: {doc.metadata.get('section')}")
            print(f"  Content: {doc.page_content[:150]}...")